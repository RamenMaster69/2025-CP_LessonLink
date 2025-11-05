import PyPDF2
import docx
import pytesseract
from PIL import Image
import os
import re  # ADD THIS IMPORT

def clean_extracted_text(text):
    """Clean up common text extraction issues"""
    if not text:
        return ""
    
    # Remove extra spaces between words
    text = re.sub(r'(\w)\s+\n\s*(\w)', r'\1 \2', text)
    
    # Fix line breaks in the middle of sentences
    text = re.sub(r'(\w+)\n\s*(\w+)', r'\1 \2', text)
    
    # Ensure proper paragraph breaks
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()

def extract_text_from_file(file_path, file_type):
    """
    Extract text from various file types
    """
    try:
        if file_type == '.txt':
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                
        elif file_type == '.pdf':
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
        elif file_type in ['.docx', '.doc']:
            doc = docx.Document(file_path)
            text = ""
            
            # First, extract all regular paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Then, extract tables if they exist
            if hasattr(doc, 'tables') and doc.tables:
                text += "\n" + "="*50 + "\n"
                text += "TABLE CONTENT:\n"
                text += "="*50 + "\n"
                
                for table_num, table in enumerate(doc.tables, 1):
                    text += f"\nTable {table_num}:\n"
                    text += "-" * 30 + "\n"
                    
                    for row_num, row in enumerate(table.rows):
                        row_data = []
                        for cell in row.cells:
                            # Get text from each cell
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_data.append(cell_text)
                            else:
                                row_data.append("")
                        
                        # Join cells with pipe separator for table format
                        text += " | ".join(row_data) + "\n"
                    
                    text += "-" * 30 + "\n"
            
        elif file_type in ['.jpg', '.jpeg', '.png']:
            # OCR for images
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
        else:
            text = "Unsupported file format"
        
        # CLEAN THE EXTRACTED TEXT BEFORE RETURNING
        return clean_extracted_text(text)
            
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def generate_lesson_prompt_from_text(text, subject, grade_level, duration):
    """Generate a prompt for AI lesson plan generation from extracted text"""
    
    # Clean the text before using it in the prompt
    cleaned_text = clean_extracted_text(text)
    
    prompt = f"""
    Create a detailed lesson plan based on the following source material:
    
    SUBJECT: {subject}
    GRADE LEVEL: {grade_level}
    DURATION: {duration} minutes
    
    SOURCE MATERIAL:
    {cleaned_text[:2000]}  # Limit text to avoid token limits
    
    Please structure the lesson plan with the following sections:
    
    ## Metadata
    **Subject:** {subject}
    **Grade Level:** {grade_level}
    **Duration:** {duration} minutes
    **Class Size:** [appropriate number]
    
    ## MELC Alignment
    **MELC Code:** [appropriate MELC code]
    **Content Standard:** [content standard]
    **Performance Standard:** [performance standard] 
    **Learning Competency:** [learning competency]
    
    ## Learning Objectives
    [3-5 specific, measurable learning objectives]
    
    ## Subject Matter
    [Topic based on the source material]
    
    ## Materials Needed
    [List of materials]
    
    ## Lesson Procedure
    
    ### A. Introduction (10-15 minutes)
    [Engaging introduction activity]
    
    ### B. Instruction/Direct Teaching (20-25 minutes)
    [Direct instruction based on source material]
    
    ### C. Guided Practice/Application (15-20 minutes)
    [Guided practice activities]
    
    ### D. Independent Practice/Evaluation (10-15 minutes)
    [Independent work or assessment]
    
    ### E. Assessment (5-10 minutes)
    [Formative assessment]
    
    ## Differentiation
    **Support for Struggling Learners:** [strategies]
    **Extension for Advanced Learners:** [activities]
    
    ## Integration
    **Values Education:** [values integration]
    **Cross-curricular:** [cross-curricular connections]
    """
    
    return prompt